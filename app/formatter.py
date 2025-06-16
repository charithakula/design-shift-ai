import re
from docx import Document
from docx.shared import Inches
import os
from io import BytesIO

def clean_text_for_docx(text: str) -> str:
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # **bold**
    text = re.sub(r'__(.*?)__', r'\1', text)      # __bold__
    text = re.sub(r'\*(.*?)\*', r'\1', text)      # *italic*
    text = re.sub(r'_(.*?)_', r'\1', text)        # _italic_
    text = text.replace('---', '')
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def remove_repeated_title(title: str, content: str) -> str:
    lines = content.splitlines()
    skip_lines = 0
    escaped_title = re.escape(title.strip())
    patterns = [
        re.compile(rf'^\s*{escaped_title}\s*$', re.IGNORECASE),
        re.compile(rf'^\s*#+\s*{escaped_title}\s*$', re.IGNORECASE),
        re.compile(rf'^\s*Title:\s*{escaped_title}\s*$', re.IGNORECASE),
        re.compile(rf'^\s*#+\s*Title:\s*{escaped_title}\s*$', re.IGNORECASE),
        re.compile(r'^\s*Content:\s*$', re.IGNORECASE),
        re.compile(r'^\s*$')
    ]
    for line in lines:
        if any(p.match(line) for p in patterns):
            skip_lines += 1
        else:
            break
    return '\n'.join(lines[skip_lines:]).lstrip('\n\r ')

def save_to_docx(sections: dict, filename: str):
    doc = Document()
    for title, content in sections.items():
        content = remove_repeated_title(title, content)
        content = clean_text_for_docx(content)
        doc.add_heading(title, level=1)
        lines = content.split('\n')
        for line in lines:
            stripped = line.strip()
            if stripped.startswith('- '):
                doc.add_paragraph(stripped[2:], style='ListBullet')
            elif re.match(r'^\d+\.\s', stripped):
                # Treat as numbered list item
                doc.add_paragraph(stripped, style='ListNumber')
            elif stripped == '':
                doc.add_paragraph('')
            else:
                doc.add_paragraph(stripped)
    doc.save(filename)

def insert_images_to_docx(filename, sections_dict, image_list):
    doc = Document()
    # Add sections first
    for title, content in sections_dict.items():
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)

    if image_list:
        doc.add_page_break()
        doc.add_heading("Converted Diagrams", level=1)

        for i, image in enumerate(image_list):
            try:
                if isinstance(image, str) and os.path.exists(image):
                    doc.add_paragraph(os.path.basename(image))
                    doc.add_picture(image, width=Inches(5.5))
                elif isinstance(image, (bytes, BytesIO)):
                    if isinstance(image, bytes):
                        image = BytesIO(image)
                    doc.add_paragraph(f"Diagram {i + 1}")
                    doc.add_picture(image, width=Inches(5.5))
            except Exception as e:
                doc.add_paragraph(f"[Failed to add image {i+1}: {e}]")

    doc.save(filename)
