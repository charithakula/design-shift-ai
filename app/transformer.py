from docx import Document

def save_to_docx(converted_dict, output_file):
    """
    Save the converted design document sections into a .docx file using python-docx.

    Args:
        converted_dict (dict): Dictionary where keys are section titles and values are converted content strings.
        output_file (str): Path to output .docx file to save.
    """
    # Create a new blank Word document
    doc = Document()

    # Loop over each converted section
    for title, content in converted_dict.items():
        # Add the section title as a heading of level 2
        doc.add_heading(title, level=2)
        # Add the converted content as a paragraph
        doc.add_paragraph(content)
        # Add an extra blank paragraph for spacing between sections
        doc.add_paragraph()

    # Save the document to the specified file path
    doc.save(output_file)


def detect_technology_from_text(text, client):
    """
    Use OpenAI GPT API to detect the primary source technology of a design document based on its text.

    Args:
        text (str): Full text extracted from the design document.
        client (OpenAI): Initialized OpenAI client instance.

    Returns:
        str: Detected technology name, or "Unknown" if detection fails or is uncertain.
    """
    # Construct a prompt instructing the model to identify the technology from the given document text
    prompt = f"""
You are an expert in software design documents.

Given the following design document content, identify the primary source technology or platform it is written for. 
If unsure, respond with "Unknown".

Content:
{text}

Answer with just the technology name.
"""

    try:
        # Call OpenAI chat completions API with a system message for context and user prompt
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You detect the technology of design documents."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=20,     # Limit response length (technology name is short)
            temperature=0      # Use deterministic output for accuracy
        )
        # Extract the detected technology from the response
        detected_tech = response.choices[0].message.content.strip()
        return detected_tech

    except Exception as e:
        # On any API error, print it and fallback to "Unknown"
        print(f"OpenAI API Error in detect_technology_from_text: {e}")
        return "Unknown"


def convert_any_to_any(section, source_tech, target_tech, client):
    """
    Convert a single design document section from source technology format to target technology format using OpenAI.

    Args:
        section (dict): Dictionary with keys "title" and "content" representing one document section.
        source_tech (str): The source technology name detected or specified.
        target_tech (str): The target technology name to convert to.
        client (OpenAI): Initialized OpenAI client instance.

    Returns:
        str: Converted text for the section or an error/warning message.
    """
    # Extract section title and content; provide defaults if missing
    title = section.get("title", "Untitled")
    content = section.get("content", "")

    # If content is empty or whitespace, return a warning string
    if not content.strip():
        return "⚠️ No content to convert."

    # Build a detailed prompt guiding the model to perform the conversion
    prompt = f"""
You are an expert in {source_tech} and {target_tech}. Convert the following {source_tech} design section into an equivalent {target_tech} format.

Title: {title}

Content:
{content}

Output:
"""

    try:
        # Call OpenAI chat completions API with instructions to convert design document sections
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": f"You convert design documents from {source_tech} to {target_tech}."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,    # Some creativity allowed to adapt the content well
            max_tokens=1500     # Allow enough tokens for detailed converted content
        )
        # Return the converted text from the response
        return response.choices[0].message.content.strip()

    except Exception as e:
        # Return an error message string on API failure
        return f"⚠️ API Error: {str(e)}"
