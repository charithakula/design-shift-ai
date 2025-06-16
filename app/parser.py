from docx import Document

def parse_docx(filepath):
    """
    Parse a .docx file into a list of sections.

    Each section is defined by:
    - A Heading 1 paragraph marking the section title.
    - All subsequent paragraphs until the next Heading 1 are the section content.

    Args:
        filepath (str): Path to the input .docx file.

    Returns:
        list of dict: List of sections, each a dict with keys:
                      'title' (str): Section title from Heading 1
                      'content' (str): Section content concatenated as text with newlines
    """
    # Load the Word document from the file
    doc = Document(filepath)
    sections = []
    current_section = {"title": None, "content": ""}  # Initialize current section holder

    # Iterate over all paragraphs in the document
    for para in doc.paragraphs:
        # Check if paragraph style is Heading 1 (section title marker)
        if para.style.name.startswith('Heading 1'):
            # If we already have a current section title, save that section before starting new
            if current_section["title"]:
                sections.append(current_section)
            # Start a new section with the current paragraph text as title
            current_section = {"title": para.text.strip(), "content": ""}
        else:
            # If inside a section, append paragraph text to current section content with a newline
            if current_section["title"]:
                current_section["content"] += para.text + "\n"

    # After loop ends, add the last section if any
    if current_section["title"]:
        sections.append(current_section)

    return sections


def save_to_docx(converted_dict, output_file):
    """
    Save converted sections into a new .docx file.

    Each key in the dictionary is treated as a section title,
    and its corresponding value as the section content.

    Args:
        converted_dict (dict): Dictionary of section titles to content strings.
        output_file (str): Path to save the output .docx file.
    """
    # Create a new Word document
    doc = Document()

    # Iterate over each section title and content
    for title, content in converted_dict.items():
        # Add the section title as a Heading 1
        doc.add_heading(title, level=1)
        # Add each line of the section content as a separate paragraph
        for line in content.strip().split("\n"):
            doc.add_paragraph(line)

    # Save the document to the specified output file
    doc.save(output_file)
